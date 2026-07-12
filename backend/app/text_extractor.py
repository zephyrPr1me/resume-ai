import io
from abc import ABC, abstractmethod
from typing import BinaryIO
from pathlib import Path
import fitz
from html.parser import HTMLParser


class _HTMLTextParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self._texts = []

    def handle_data(self, data):
        self._texts.append(data)

    def get_text(self):
        return "\n".join([t.strip() for t in self._texts if t.strip()])


class BaseTextExtractor(ABC):
    def __init__(self, file_stream: BinaryIO, filename: str):
        self.file_stream = file_stream
        self.filename = filename
        self._validate_stream()

    def _validate_stream(self):
        if not isinstance(self.file_stream, io.IOBase):
            raise TypeError("Expected a file-like object (e.g. io.BytesIO)")

        self.file_stream.seek(0, io.SEEK_END)
        size = self.file_stream.tell()
        self.file_stream.seek(0)

        if size > 50 * 1024 * 1024:  # Limit to 50 MB
            raise ValueError("The file is too large to process in memory (>50MB)")

    @abstractmethod
    def extract(self) -> str:
        pass


class PdfExtractor(BaseTextExtractor):
    def extract(self) -> str:
        text = []
        with fitz.open(stream=self.file_stream, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text())
        return "\n".join(text)


class TxtExtractor(BaseTextExtractor):
    def extract(self) -> str:
        self.file_stream.seek(0)
        return self.file_stream.read().decode("utf-8")


class MdExtractor(TxtExtractor):
    """Markdown extractor — same as text for now."""

    pass


class HtmlExtractor(BaseTextExtractor):
    def extract(self) -> str:
        self.file_stream.seek(0)
        raw = self.file_stream.read().decode("utf-8", errors="ignore")
        parser = _HTMLTextParser()
        parser.feed(raw)
        return parser.get_text()


class DocxExtractor(BaseTextExtractor):
    def extract(self) -> str:
        try:
            from docx import Document
        except Exception as e:
            raise ValueError(
                "python-docx is required to extract .docx files. Install with `pip install python-docx`"
            ) from e

        # python-docx expects a path or a file-like object; BytesIO works
        self.file_stream.seek(0)
        try:
            doc = Document(self.file_stream)
        except Exception:
            # Some docx files may need to be written to a temp file — try fallback
            from tempfile import NamedTemporaryFile

            self.file_stream.seek(0)
            with NamedTemporaryFile(delete=True, suffix=".docx") as tmp:
                tmp.write(self.file_stream.read())
                tmp.flush()
                doc = Document(tmp.name)

        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return "\n".join(paragraphs)


def get_stream_extractor(file_stream: BinaryIO, filename: str) -> BaseTextExtractor:
    ext = Path(filename).suffix.lower()

    extractors_map = {
        ".txt": TxtExtractor,
        ".pdf": PdfExtractor,
        ".md": MdExtractor,
        ".markdown": MdExtractor,
        ".html": HtmlExtractor,
        ".htm": HtmlExtractor,
        ".docx": DocxExtractor,
    }

    extractor_class = extractors_map.get(ext)
    if not extractor_class:
        raise ValueError(f"Unsupported format: {ext}")

    return extractor_class(file_stream, filename)
